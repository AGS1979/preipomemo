import streamlit as st
import os
import re
import fitz  # PyMuPDF
import faiss
import json
import requests
import markdown
import numpy as np
from datetime import datetime
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt, Inches
from sentence_transformers import SentenceTransformer
from collections import defaultdict

# ========== CONFIG ==========
DEEPSEEK_API_KEY = st.secrets["deepseek"]["api_key"]
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"
CHUNK_SIZE = 50


# ========== CORE UTILITIES ==========

def clean_markdown(text):
    text = re.sub(r'#+\s*', '', text)  # remove markdown headers
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # bold markdown
    text = re.sub(r'\*(.*?)\*', r'\1', text)  # italic markdown
    text = re.sub(r'_+', '', text)  # underscore artifacts
    text = re.sub(r'\n{3,}', '\n\n', text)  # excess newlines
    text = re.sub(r'^[-*•]+\s+', '', text, flags=re.MULTILINE)  # bullets
    text = re.sub(r'Section\s\d+[:.]?', '', text, flags=re.IGNORECASE)  # e.g., "Section 2:"
    text = re.sub(r'(Next|Previous) section:.*', '', text, flags=re.IGNORECASE)
    text = re.sub(r'This section .*?(focuses on|explores|explains).*?\.', '', text, flags=re.IGNORECASE)
    return text.strip()


def extract_text_by_page(pdf_path):
    doc = fitz.open(pdf_path)
    return [page.get_text() for page in doc], len(doc)

def get_relevant_pages_chunked(text_by_page, user_query):
    total_pages = len(text_by_page)
    relevant_pages = set()

    for start in range(0, total_pages, CHUNK_SIZE):
        end = min(start + CHUNK_SIZE, total_pages)
        chunk_pages = text_by_page[start:end]

        prompt = (
            "Below are texts extracted from pages of a PDF. Identify only the page numbers (starting from 1) relevant to this query:\n"
            f"Query: {user_query}\n\n"
        )
        for i, text in enumerate(chunk_pages):
            snippet = text[:1000].replace('\n', ' ')
            prompt += f"\nPage {start + i + 1}: {snippet}\n"

        messages = [
            {"role": "system", "content": "You are an expert document analyst."},
            {"role": "user", "content": prompt},
        ]

        response = requests.post(
            DEEPSEEK_API_URL,
            headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}"},
            json={"model": "deepseek-chat", "messages": messages}
        )
        response.raise_for_status()

        reply = response.json()['choices'][0]['message']['content']
        matches = re.findall(r'\d+', reply)
        for m in matches:
            num = int(m)
            if 1 <= num <= total_pages:
                relevant_pages.add(num)

    return sorted(relevant_pages)

def extract_selected_pages_text(original_path, pages_to_keep):
    doc = fitz.open(original_path)
    return "\n".join(doc[p - 1].get_text() for p in pages_to_keep).strip()

def extract_company_name(text):
    prompt = (
        "Extract only the legal name of the company from the following IPO or DRHP text. "
        "Return only the company name, nothing else.\n\n"
        f"{text[:3000]}"
    )
    messages = [
        {"role": "system", "content": "You are an expert in IPO documents."},
        {"role": "user", "content": prompt},
    ]
    response = requests.post(
        DEEPSEEK_API_URL,
        headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}"},
        json={"model": "deepseek-chat", "messages": messages}
    )
    response.raise_for_status()
    return response.json()['choices'][0]['message']['content'].strip()

def generate_memo_sections(filtered_text, custom_notes=""):
    section_titles = [
        "1. IPO Offer Details", "2. Company Overview", "3. Industry Overview and Outlook",
        "4. Business Model", "5. Financial Highlights",
        "6. Guidance and Outlook on future financial performance",
        "7. Peer Comparison and Competitors", "8. Risks", "9. Investment Highlights"
    ]

    sections = {}
    for title in section_titles:
        prompt = (
            f"You are writing a professional pre-IPO investment memo section titled: {title[3:]}. "
            "Please generate ~500 words of clean, structured, analytical prose suitable for institutional investors. "
            "Do not mention this is a memo. Avoid starting with the section title, and avoid phrases like 'In this section' or 'previously discussed'. "
            "Strictly avoid markdown (no asterisks, hashes, underscores). Use plain text only.\n\n"
        )

        if custom_notes:
            prompt += f"Focus on this angle: {custom_notes.strip()}\n\n"

        prompt += f"Relevant DRHP Text:\n{filtered_text[:16000]}"


        messages = [
            {"role": "system", "content": "You are an expert financial analyst."},
            {"role": "user", "content": prompt},
        ]
        response = requests.post(
            DEEPSEEK_API_URL,
            headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}"},
            json={"model": "deepseek-chat", "messages": messages}
        )
        response.raise_for_status()
        raw_content = response.json()['choices'][0]['message']['content']
        cleaned = clean_markdown(raw_content)
        cleaned = re.sub(rf"^{re.escape(title[3:])}[\s:—-]*", "", cleaned, flags=re.IGNORECASE)
        sections[title] = cleaned

    return sections

def save_sections_to_word(sections_dict, company_name="Company", output_dir="documents"):
    os.makedirs(output_dir, exist_ok=True)
    filename = f"{company_name.replace(' ', '_')}_PreIPO_Memo_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    full_path = os.path.join(output_dir, filename)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Aptos Display'
    style.font.size = Pt(11)

    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"{company_name} Pre-IPO Investment Memo")
    title_run.font.name = 'Aptos Display'
    title_run.font.size = Pt(20)
    title_run.bold = True
    doc.add_paragraph()

    for title, body in sections_dict.items():
        heading = doc.add_paragraph()
        run = heading.add_run(title)
        run.bold = True
        run.font.name = 'Aptos Display'
        run.font.size = Pt(14)

        for para in body.strip().split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.add_paragraph()

    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.gutter = Inches(0)

    doc.save(full_path)
    return full_path


# ========== PIPELINE WRAPPER ==========

def run_pipeline(pdf_path, custom_focus="", output_dir="documents"):
    text_by_page, _ = extract_text_by_page(pdf_path)

    default_query = (
        "Extract only those pages that contain information useful for writing a pre-IPO investment memo. "
        "This includes sections on 'Management’s Discussion and Analysis of Financial Condition and Results of Operations', "
        "'Financial Highlights', 'Risk Factors', 'Business Overview', and 'Industry Overview'. Exclude all other pages."
    )

    pages_to_keep = get_relevant_pages_chunked(text_by_page, default_query)
    if not pages_to_keep:
        raise ValueError("No relevant pages found.")

    filtered_text = extract_selected_pages_text(pdf_path, pages_to_keep)
    if not filtered_text.strip():
        raise ValueError("Filtered text is empty.")

    company_name = extract_company_name(filtered_text)
    sections_dict = generate_memo_sections(filtered_text, custom_focus)
    return save_sections_to_word(sections_dict, company_name=company_name, output_dir=output_dir)


# ========== PDF Q&A ENGINE ==========

class PDFQueryEngine:
    def __init__(self, api_key=DEEPSEEK_API_KEY, model_name="all-MiniLM-L6-v2"):
        if not api_key:
            raise ValueError("❌ DEEPSEEK_API_KEY is not set.")
        self.api_key = api_key
        self.embedder = SentenceTransformer(model_name)

    def extract_text_from_pdf(self, path):
        reader = PdfReader(path)
        return [(i + 1, page.extract_text().strip()) for i, page in enumerate(reader.pages) if page.extract_text()]

    def embed_texts(self, texts):
        return np.array(self.embedder.encode(texts, convert_to_numpy=True))

    def build_faiss_index(self, embs):
        dim = embs.shape[1]
        index = faiss.IndexFlatL2(dim)
        index.add(embs)
        return index

    def query_deepseek(self, context_chunks, query):
        messages = [{"role": "system", "content": "Answer questions based on the context provided."}]
        for page, text in context_chunks:
            messages.append({"role": "user", "content": f"[Page {page}] {text}"})
        messages.append({"role": "user", "content": query})

        response = requests.post(
            DEEPSEEK_API_URL,
            headers={"Authorization": f"Bearer {self.api_key}", "Content-Type": "application/json"},
            json={"model": "deepseek-chat", "messages": messages, "temperature": 0.3}
        )
        response.raise_for_status()
        md_response = response.json()["choices"][0]["message"]["content"]
        return markdown.markdown(md_response)

    def answer_query(self, pdf_path, query, top_k=3):
        chunks = self.extract_text_from_pdf(pdf_path)
        if not chunks:
            raise ValueError("No text extracted from PDF.")

        pages, texts = zip(*chunks)
        embs = self.embed_texts(texts)
        index = self.build_faiss_index(embs)

        q_emb = self.embed_texts([query])
        _, I = index.search(q_emb, k=top_k)
        selected = [(pages[i], texts[i]) for i in I[0]]
        answer = self.query_deepseek(selected, query)
        cited_pages = [pages[i] for i in I[0]]
        return answer, cited_pages



def extract_raw_text(docx_path):
    doc = Document(docx_path)
    return "\n".join(para.text.strip() for para in doc.paragraphs if para.text.strip())

def call_deepseek_summary(text, company_name):
    prompt = f"""
You are an investment analyst tasked with converting the following pre-IPO memo into a concise infographic-ready summary.

Summarize the key points for each of these sections:
1. IPO Offer Details
2. Company Overview
3. Industry Overview and Outlook
4. Business Model
5. Financial Highlights
6. Guidance and Outlook on future financial performance
7. Peer Comparison and Competitors
8. Risks
9. Investment Highlights

Each section should contain 3–5 bullet points maximum. Use crisp, bullet-style formatting (no paragraphs). Keep each bullet point under 30 words.

Company: {company_name}
Memo:
{text}
"""

    headers = {
        "Authorization": f"Bearer {os.getenv('DEEPSEEK_API_KEY')}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": "deepseek-chat",
        "messages": [
            {"role": "system", "content": "You are a helpful analyst."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.3
    }

    response = requests.post("https://api.deepseek.com/v1/chat/completions", headers=headers, json=payload)
    response.raise_for_status()
    return response.json()["choices"][0]["message"]["content"]

def bold_labels(text):
    return re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)

def parse_deepseek_response(summary_text):
    sections = defaultdict(list)
    current_section = None
    lines = summary_text.splitlines()
    for line in lines:
        line = line.strip()
        section_match = re.match(r"^#+\s+\*+\d+\.\s+(.*?)\*+\s*$", line)
        if section_match:
            current_section = section_match.group(1).strip()
            continue
        simple_header = re.match(r"^#+\s+(.*)", line)
        if simple_header:
            current_section = simple_header.group(1).strip()
            continue
        bullet_match = re.match(r"^- (.+)", line)
        if bullet_match and current_section:
            bullet = bold_labels(bullet_match.group(1).strip())
            sections[current_section].append(bullet)
    return dict(sections)

from jinja2 import Template

def generate_infographic_html(docx_path, company_name, base_template_path="base_infographic.html"):
    raw_text = extract_raw_text(docx_path)
    summary = call_deepseek_summary(raw_text, company_name)
    sections = parse_deepseek_response(summary)

    with open(base_template_path, "r", encoding="utf-8") as f:
        html_template = f.read()

    template = Template(html_template)
    rendered_html = template.render(company_name=company_name, sections=sections)

    return rendered_html
